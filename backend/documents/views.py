import io
import pdfplumber
import pytesseract
from PIL import Image
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.db import models
from rest_framework import generics, viewsets, status, filters
from rest_framework.response import Response
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from django.contrib.auth.models import User

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from rest_framework.views import APIView

from core.mixins import StandardResponseMixin
from users.models import Usuarios
from users.serializers import UsuariosSerializer

from .models import (
    DocumentoSubido, 
    CampoDisponible, 
    PlantillaDocumento, 
    CampoPlantilla, 
    DocumentoGenerado,
    PlantillaFavorita,
    TipoPlantillaDocumento,
    CategoriaPlantillaDocumento,
    PlantillaCompartida,
    ClasificacionPlantillaGeneral,
    PlantillaGeneral,
    PlantillaGeneralCompartida,
)
from .serializers import (
    DocumentoSubidoSerializer,
    CampoDisponibleSerializer,
    PlantillaDocumentoSerializer,
    DocumentoGeneradoSerializer,
    CrearPlantillaSerializer,
    GenerarDocumentoSerializer,
    TipoPlantillaDocumentoSerializer,
    CategoriaPlantillaDocumentoSerializer,
    PlantillaCompartidaSerializer,
    PlantillaFavoritaSerializer,
    ClasificacionPlantillaGeneralSerializer,
    PlantillaGeneralSerializer,
    PlantillaGeneralCompartidaSerializer,
    FileUploadSerializer,
)


class DocumentoSubidoViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = DocumentoSubido.objects.all()
    serializer_class = DocumentoSubidoSerializer
    parser_classes = (MultiPartParser, FormParser)

    def list(self, request, *args, **kwargs):
        """Listar documentos subidos con filtrado por roles"""
        try:
            user = request.user
            
            # Si es superuser o staff, puede ver todos los documentos
            if user.is_superuser or user.is_staff:
                queryset = DocumentoSubido.objects.all()
            
            # Si pertenece al grupo 'Admin', puede ver documentos de usuarios de su empresa
            elif user.groups.filter(name='Admin').exists():
                if user.empresa:
                    # Obtener todos los usuarios de la misma empresa
                    usuarios_empresa = Usuarios.objects.filter(empresa=user.empresa)
                    queryset = DocumentoSubido.objects.filter(usuario__in=usuarios_empresa)
                else:
                    # Si no tiene empresa asignada, solo ve sus documentos
                    queryset = DocumentoSubido.objects.filter(usuario=user)
            
            # Usuario común: solo ve sus propios documentos
            else:
                queryset = DocumentoSubido.objects.filter(usuario=user)
            
            # Ordenar por fecha de subida descendente
            queryset = queryset.order_by('-fecha_subida')
            
            serializer = self.get_serializer(queryset, many=True)
            return self.success_response(
                data=serializer.data,
                message="Documentos subidos obtenidos exitosamente",
                code="documents_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener documentos subidos",
                code="documents_error",
                http_status=500
            )
    
    def create(self, request, *args, **kwargs):
        """Crear documento subido con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Documento subido exitosamente",
            code="document_created",
            error_message="Error al subir documento",
            error_code="document_create_error"
        )

    @action(detail=False, methods=['post'])
    def subir_documento(self, request):
        """Subir documento y extraer texto"""
        try:
            archivo = request.FILES.get('archivo')
            if not archivo:
                return self.error_response(
                    errors="No se proporcionó archivo",
                    message="Archivo requerido",
                    code="missing_file",
                    http_status=400
                )

            # Determinar tipo de archivo
            nombre_archivo = archivo.name.lower()
            if nombre_archivo.endswith('.pdf'):
                tipo = 'pdf'
                #ESTO ES PARA PRUEBA REUNION DEL 22
                return self.error_response(
                    errors="Solo se permiten archivos docx",
                    message="Archivo requerido",
                    code="missing_file",
                    http_status=400
                )
            elif nombre_archivo.endswith(('.jpg', '.jpeg', '.png', '.gif')):
                tipo = 'imagen'
                #ESTO ES PARA PRUEBA REUNION DEL 22
                return self.error_response(
                    errors="Solo se permiten archivos docx",
                    message="Archivo requerido",
                    code="missing_file",
                    http_status=400
                )
            elif nombre_archivo.endswith('.docx'):
                tipo = 'word'
            else:
                tipo = 'texto'
                #ESTO ES PARA PRUEBA REUNION DEL 22
                return self.error_response(
                    errors="Solo se permiten archivos docx",
                    message="Archivo requerido",
                    code="missing_file",
                    http_status=400
                )
            
            if request.user.is_staff or request.user.is_superuser:
                # Administradores y staff guardan en carpeta admin
                ruta_archivo = f'documentos/admin/{request.user.username}/{archivo.name}'
            else:
                # Usuarios regulares guardan en carpeta empresa/usuario
                if hasattr(request.user, 'empresa') and request.user.empresa:
                    ruta_archivo = f'documentos/{request.user.empresa.id}/{request.user.username}/{archivo.name}'
                else:
                    # Fallback si el usuario no tiene empresa asignada
                    ruta_archivo = f'documentos/sin_empresa/{request.user.username}/{archivo.name}'

            #ESTO ES PARA PRUEBA REUNION DEL 22
            #ruta_completa = default_storage.save(ruta_archivo, ContentFile(archivo.read()))

            # Extraer texto según el tipo
            texto_extraido = ""
            if tipo == 'pdf':
                texto_extraido = self._extraer_texto_pdf(archivo)
            elif tipo == 'imagen':
                texto_extraido = self._extraer_texto_imagen(archivo)
            elif tipo == 'word':
                # Guardar archivo temporalmente para python-docx
                archivo.seek(0)
                doc = docx.Document(archivo)
                texto_extraido = self._extraer_texto_docx(doc)
            else:
                archivo.seek(0)
                texto_extraido = archivo.read().decode('utf-8')
                
            #print("texto extraido: ", texto_extraido)

            # Verificar autenticación
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            
            # Crear registro en base de datos
            # COMENTAR PARA QUE NO GUARDE EL REGISTRO SEGUN REUNION DEL 22
            '''documento = DocumentoSubido.objects.create(
                usuario=request.user,
                nombre_original=archivo.name,
                tipo=tipo,
                archivo_url='', #ruta_completa
                html=texto_extraido
            )'''

            # Preparar datos de respuesta
            data = {
                'id': '', #documento.id
                'html': texto_extraido,
                'tipo': tipo,
                'nombre_original': archivo.name,
                'archivo_url': '', #ruta_completa
                'fecha_subida': '' #documento.fecha_subida
            }

            return self.success_response(
                data=data,
                message="Documento subido y procesado exitosamente",
                code="document_uploaded",
                http_status=201
            )

        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al procesar el documento",
                code="document_processing_error",
                http_status=500
            )

    def _extraer_texto_pdf(self, archivo):
        """Extraer texto de PDF usando pdfplumber"""
        try:
            pdf = pdfplumber.open(archivo)
            html = ""
            
            for page in pdf.pages:
                words = page.extract_words()
                page_width = page.width
                
                # Agrupa palabras por línea (top)
                lines = {}
                for word in words:
                    line_key = round(word['top'])
                    if line_key not in lines:
                        lines[line_key] = []
                    lines[line_key].append(word)
                
                # Procesar cada línea
                for line_words in lines.values():
                    if not line_words:
                        continue
                        
                    line_words.sort(key=lambda w: w['x0'])
                    
                    # Calcular posiciones para detectar centrado
                    first_word = line_words[0]
                    last_word = line_words[-1]
                    line_start = first_word['x0']
                    line_end = last_word['x1']
                    line_width = line_end - line_start
                    
                    # Detectar si es centrado: debe tener espacio significativo tanto al inicio como al final
                    margin_left = line_start
                    margin_right = page_width - line_end
                    min_margin = page_width * 0.15  # Al menos 15% de margen en cada lado
                    
                    # Una línea está centrada si:
                    # 1. Tiene márgenes significativos en ambos lados
                    # 2. Los márgenes son relativamente similares (diferencia < 20% del ancho de página)
                    # 3. No ocupa más del 70% del ancho de la página
                    is_centered = (
                        margin_left > min_margin and 
                        margin_right > min_margin and
                        abs(margin_left - margin_right) < (page_width * 0.2) and
                        line_width < (page_width * 0.7)
                    )
                    
                    # Calcular indentación para líneas no centradas
                    indent_level = 0
                    if not is_centered and margin_left > 20:
                        indent_level = max(0, int(margin_left / 25))
                    
                    # Construir el texto de la línea
                    line_text = ""
                    prev_x1 = None
                    
                    for idx, word in enumerate(line_words):
                        if idx == 0:
                            # Primera palabra de la línea
                            line_text += word['text']
                        else:
                            gap = word['x0'] - prev_x1
                            
                            if gap > 40:  # Espacio muy grande - posible tabulación
                                spaces = "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"  # 5 espacios no separables
                            elif gap > 20:  # Espacio grande
                                spaces = "&nbsp;&nbsp;"  # 2 espacios no separables
                            else:  # Espacio normal
                                spaces = " "
                            
                            line_text += spaces + word['text']
                        
                        prev_x1 = word['x1']
                    
                    # Aplicar estilos según el tipo de línea
                    if is_centered:
                        # Línea centrada - probablemente un título
                        html += f"<p style='text-align: center; font-weight: bold; margin: 12px 0; font-size: 14px;'>{line_text.strip()}</p>\n"
                    else:
                        # Línea normal - aplicar justificación y sangría
                        indent_style = f"margin-left: {indent_level * 15}px; " if indent_level > 0 else ""
                        
                        # Detectar si es una línea larga que debería justificarse
                        # (más del 60% del ancho disponible y más de 6 palabras)
                        should_justify = (
                            line_width > (page_width * 0.6) and 
                            len(line_words) > 6 and 
                            not line_text.strip().endswith(':')  # No justificar líneas que terminan en :
                        )
                        
                        text_align = "text-align: justify; " if should_justify else "text-align: left; "
                        
                        html += f"<p style='{indent_style}{text_align}margin: 6px 0; line-height: 1.4;'>{line_text.strip()}</p>\n"
            
            pdf.close()
            
            # Envolver todo en un contenedor
            final_html = f"""
            <div style='font-family: "Times New Roman", serif; line-height: 1.5; max-width: 800px; margin: 0 auto; padding: 20px; font-size: 13px;'>
                {html}
            </div>
            """

            #print(final_html)
            
            return final_html.strip()
            
        except Exception as e:
            return f"Error al extraer texto del PDF: {str(e)}"

    def _extraer_texto_imagen(self, archivo):
        """Extraer texto de imagen usando OCR manteniendo formato visual"""
        try:
            img = Image.open(archivo)
            
            # Lista de idiomas a probar en orden de preferencia
            idiomas = ['eng', 'spa', 'spa+eng']
            
            for idioma in idiomas:
                try:
                    # Usar pytesseract para obtener datos de cada palabra con posiciones
                    ocr_data = pytesseract.image_to_data(img, lang=idioma, output_type=pytesseract.Output.DICT)
                    break
                except pytesseract.TesseractError:
                    continue
            else:
                # Si ningún idioma funcionó, usar configuración por defecto
                ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            
            n = len(ocr_data['level'])
            paragraphs = {}
            
            # Agrupar palabras por párrafos y líneas
            for i in range(n):
                if int(ocr_data['conf'][i]) > 0 and ocr_data['text'][i].strip():
                    block_num = ocr_data['block_num'][i]
                    par_num = ocr_data['par_num'][i]
                    line_num = ocr_data['line_num'][i]
                    
                    # Usar block_num y par_num para agrupar párrafos
                    par_key = (block_num, par_num)
                    line_key = (block_num, par_num, line_num)
                    
                    if par_key not in paragraphs:
                        paragraphs[par_key] = {}
                    
                    if line_key not in paragraphs[par_key]:
                        paragraphs[par_key][line_key] = []
                    
                    paragraphs[par_key][line_key].append({
                        'text': ocr_data['text'][i],
                        'left': ocr_data['left'][i],
                        'top': ocr_data['top'][i],
                        'width': ocr_data['width'][i],
                        'height': ocr_data['height'][i]
                    })
            
            html = ""
            img_width = img.width
            
            # Procesar párrafos ordenados
            for par_key in sorted(paragraphs.keys()):
                paragraph_lines = paragraphs[par_key]
                paragraph_html = ""
                
                # Procesar líneas del párrafo
                for line_key in sorted(paragraph_lines.keys(), key=lambda k: k[2]):  # Ordenar por line_num
                    line = paragraph_lines[line_key]
                    if not line:
                        continue
                    
                    # Ordenar palabras por posición horizontal
                    line.sort(key=lambda w: w['left'])
                    
                    # Construir texto de la línea con espaciado preciso
                    line_text = ""
                    first_word = line[0]
                    last_word = line[-1]
                    
                    # Calcular posiciones para detectar centrado
                    line_start = first_word['left']
                    line_end = last_word['left'] + last_word['width']
                    line_width = line_end - line_start
                    
                    # Detectar si es centrado: debe tener espacio significativo tanto al inicio como al final
                    margin_left = line_start
                    margin_right = img_width - line_end
                    min_margin = img_width * 0.15  # Al menos 15% de margen en cada lado
                    
                    # Una línea está centrada si:
                    # 1. Tiene márgenes significativos en ambos lados
                    # 2. Los márgenes son relativamente similares (diferencia < 20% del ancho de imagen)
                    # 3. No ocupa más del 70% del ancho de la imagen
                    is_centered = (
                        margin_left > min_margin and 
                        margin_right > min_margin and
                        abs(margin_left - margin_right) < (img_width * 0.2) and
                        line_width < (img_width * 0.7)
                    )
                    
                    # Calcular indentación para líneas no centradas
                    indent_level = 0
                    if not is_centered and margin_left > 20:
                        indent_level = max(0, int(margin_left / 25))
                    
                    # Construir el texto de la línea
                    line_text += first_word['text']
                    prev_right = first_word['left'] + first_word['width']
                    
                    # Agregar resto de palabras con espaciado apropiado
                    for word in line[1:]:
                        gap = word['left'] - prev_right
                        
                        if gap > 40:  # Espacio muy grande - posible tabulación
                            spaces = "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"  # 5 espacios no separables
                        elif gap > 20:  # Espacio grande
                            spaces = "&nbsp;&nbsp;"  # 2 espacios no separables
                        else:  # Espacio normal
                            spaces = " "
                        
                        line_text += spaces + word['text']
                        prev_right = word['left'] + word['width']
                    
                    # Aplicar estilos según el tipo de línea
                    if is_centered:
                        # Línea centrada - probablemente un título
                        paragraph_html += f"<p style='text-align: center; font-weight: bold; margin: 12px 0; font-size: 14px;'>{line_text.strip()}</p>\n"
                    else:
                        # Línea normal - aplicar justificación y sangría
                        indent_style = f"margin-left: {indent_level * 15}px; " if indent_level > 0 else ""
                        
                        # Detectar si es una línea larga que debería justificarse
                        # (más del 60% del ancho disponible y más de 6 palabras)
                        should_justify = (
                            line_width > (img_width * 0.6) and 
                            len(line) > 6 and 
                            not line_text.strip().endswith(':')  # No justificar líneas que terminan en :
                        )
                        
                        text_align = "text-align: justify; " if should_justify else "text-align: left; "
                        
                        paragraph_html += f"<p style='{indent_style}{text_align}margin: 6px 0; line-height: 1.4;'>{line_text.strip()}</p>\n"
                
                # Agregar el párrafo completo
                if paragraph_html.strip():
                    html += f"<div style='margin-bottom: 12px;'>\n{paragraph_html}</div>\n"
            
            # Envolver todo en un contenedor
            final_html = f"""
            <div style='font-family: "Times New Roman", serif; line-height: 1.5; max-width: 800px; margin: 0 auto; padding: 20px; font-size: 13px;'>
                {html}
            </div>
            """
            
            return final_html.strip()
            
        except Exception as e:
            return f"Error al extraer texto de imagen: {str(e)}"
            
            

    def _extraer_texto_docx(self, doc):
        """Extrae texto de un documento DOCX y lo convierte a HTML manteniendo el formato"""
        html = self._get_base_css_styles()
        in_list = False
        
        for block in self._iter_block_items(doc):
            if isinstance(block, Table):
                html += self._process_table(block)
            elif isinstance(block, docx.text.paragraph.Paragraph):
                paragraph_html, in_list = self._process_paragraph(block, in_list)
                html += paragraph_html
        
        if in_list:
            html += "</ul>"
        
        print("html de word: ", html)
        return html
    
    def _get_base_css_styles(self):
        """Retorna los estilos CSS base para el documento HTML generado"""
        return (
            "<style>"
            "body{font-family:Arial,sans-serif;margin:0;padding:20px;line-height:1.2;}"
            "p{margin:0;padding:0;}"
            "table{border-collapse:collapse;}"
            "td{vertical-align:top;}"
            "</style>"
        )
    
    def _process_table(self, table):
        """Procesa una tabla DOCX y retorna su representación HTML"""
        table_html = "<table style='border-collapse:collapse;margin:10px 0;width:100%;border:1px solid #ccc;'>"
        
        for row in table.rows:
            table_html += "<tr>"
            for cell in row.cells:
                # Procesar todos los párrafos de la celda
                cell_content = ""
                if cell.paragraphs:
                    for i, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():  # Solo procesar párrafos con contenido
                            # Obtener estilos del párrafo
                            css_styles = self._get_paragraph_styles(paragraph)
                            paragraph_content = self._process_paragraph_runs(paragraph)
                            
                            # Aplicar estilos si existen
                            if css_styles:
                                cell_content += f'<div style="{css_styles}">{paragraph_content}</div>'
                            else:
                                cell_content += f'<div>{paragraph_content}</div>'
                        elif i == 0 and not cell_content:  # Celda vacía
                            cell_content = "&nbsp;"
                else:
                    cell_content = "&nbsp;"  # Celda sin párrafos
                
                # Aplicar estilos de celda
                cell_style = "padding:8px;vertical-align:top;border:1px solid #ddd;"
                table_html += f"<td style='{cell_style}'>{cell_content}</td>"
            table_html += "</tr>"
        
        table_html += "</table>"
        return table_html
    
    def _process_paragraph(self, paragraph, is_currently_in_list):
        """Procesa un párrafo DOCX y retorna su HTML junto con el estado actualizado de lista"""
        paragraph_style = paragraph.style.name.lower()
        formatted_content = self._process_paragraph_runs(paragraph)
        
        # Obtener todos los estilos CSS del párrafo
        css_styles = self._get_paragraph_styles(paragraph)
        style_attribute = f'style="{css_styles}"' if css_styles else ""
        
        # Determinar si es un elemento de lista
        if self._is_list_paragraph(paragraph):
            paragraph_html, updated_list_state = self._handle_list_paragraph(
                formatted_content, style_attribute, is_currently_in_list
            )
        else:
            paragraph_html, updated_list_state = self._handle_regular_paragraph(
                formatted_content, style_attribute, paragraph_style, is_currently_in_list
            )
        
        return paragraph_html, updated_list_state
    
    def _handle_list_paragraph(self, content, style_attr, in_list):
        """Maneja párrafos que son elementos de lista"""
        html_output = ""
        
        if not in_list:
            html_output = "<ul style='margin:0;padding-left:20px;'>"
            in_list = True
        
        html_output += f"<li {style_attr}>{content}</li>"
        return html_output, in_list
    
    def _handle_regular_paragraph(self, content, style_attr, paragraph_style, in_list):
        """Maneja párrafos regulares (no de lista)"""
        html_output = ""
        
        # Cerrar lista si estábamos en una
        if in_list:
            html_output += "</ul>"
            in_list = False
        
        # Determinar el tipo de elemento HTML según el estilo
        if paragraph_style.startswith("heading"):
            html_output += f"<h2 {style_attr}>{content}</h2>"
        else:
            html_output += f"<p {style_attr}>{content}</p>"
        
        return html_output, in_list
    
    def _get_paragraph_styles(self, paragraph):
        """Extrae y combina todos los estilos de un párrafo"""
        styles = []

        # Alineación
        alignment = self._get_alignment_style(paragraph)
        if alignment:
            styles.append(alignment)
        
        # Sangrías
        indentation = self._get_indentation_styles(paragraph)
        if indentation:
            styles.append(indentation)
        
        # Espaciado
        spacing = self._get_spacing_styles(paragraph)
        if spacing:
            styles.append(spacing)
        
        # Interlineado
        line_spacing = self._get_line_spacing_style(paragraph)
        if line_spacing:
            styles.append(line_spacing)
        
        return "".join(styles)
    
    def _get_alignment_style(self, paragraph):
        """Obtiene el estilo de alineación del párrafo"""
        try:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.CENTER: "text-align:center;",
                WD_ALIGN_PARAGRAPH.RIGHT: "text-align:right;",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "text-align:justify;",
                WD_ALIGN_PARAGRAPH.LEFT: "text-align:left;"
            }
            return alignment_map.get(paragraph.alignment, "")
        except Exception as e:
            # Manejar casos donde el valor de alineación no tiene mapeo XML válido
            # como 'start' que puede aparecer en algunos documentos DOCX
            #print(f"Warning: Alignment value not supported: {e}")
            return ""
    
    def _get_indentation_styles(self, paragraph):
        """Obtiene los estilos de sangría del párrafo"""
        styles = []
        fmt = paragraph.paragraph_format
        
        # Sangría de primera línea
        if fmt.first_line_indent and fmt.first_line_indent.pt != 0:
            styles.append(f"text-indent:{fmt.first_line_indent.pt:.1f}pt;")
        
        # Sangría izquierda
        if fmt.left_indent and fmt.left_indent.pt != 0:
            styles.append(f"margin-left:{fmt.left_indent.pt:.1f}pt;")
        
        # Sangría derecha
        if fmt.right_indent and fmt.right_indent.pt != 0:
            styles.append(f"margin-right:{fmt.right_indent.pt:.1f}pt;")
        
        return "".join(styles)
    
    def _get_spacing_styles(self, paragraph):
        """Obtiene los estilos de espaciado del párrafo"""
        styles = []
        fmt = paragraph.paragraph_format
        
        if fmt.space_before and fmt.space_before.pt > 0:
            styles.append(f"margin-top:{fmt.space_before.pt:.1f}pt;")
        
        if fmt.space_after and fmt.space_after.pt > 0:
            styles.append(f"margin-bottom:{fmt.space_after.pt:.1f}pt;")
        
        return "".join(styles)
    
    def _get_line_spacing_style(self, paragraph):
        """Obtiene el estilo de interlineado del párrafo"""

        fmt = paragraph.paragraph_format
        if not fmt.line_spacing:
            return ""
        #print("TAMAÑO DE INTERLINEADO:", fmt.line_spacing)
        if fmt.line_spacing_rule == 1:  # Múltiple
            return f"line-height:{fmt.line_spacing:.1f};"
        elif fmt.line_spacing_rule in [0, 2]:  # Exacto o Mínimo
            return f"line-height:{fmt.line_spacing.pt:.1f}pt;"
        
        return ""
    
    def _process_paragraph_runs(self, paragraph):
        """Procesa los runs de un párrafo para mantener el formato de negrita y otros estilos"""
        if not paragraph:
            return ""
        
        formatted_text = ""
        for run in paragraph.runs:
            text = self._process_run_text(run)
            formatted_text += text
        
        return formatted_text
    
    def _process_run_text(self, run):
        """Procesa un run individual aplicando todos los formatos necesarios"""
        text = run.text
        
        # Preservar espacios múltiples
        text = self._preserve_multiple_spaces(text)
        
        # Aplicar formatos de texto
        text = self._apply_text_formatting(run, text)
        
        # Aplicar estilos de fuente
        text = self._apply_font_styles(run, text)
        
        return text
    
    def _preserve_multiple_spaces(self, text):
        """Convierte espacios múltiples consecutivos a entidades HTML"""
        import re
        return re.sub(r'  +', lambda m: '&nbsp;' * len(m.group()), text)
    
    def _apply_text_formatting(self, run, text):
        """Aplica formato de texto básico (negrita, cursiva, subrayado)"""
        if run.bold:
            text = f"<strong>{text}</strong>"
        
        if run.italic:
            text = f"<em>{text}</em>"
        
        if run.underline:
            text = f"<u>{text}</u>"
        
        return text
    
    def _apply_font_styles(self, run, text):
        """Aplica estilos de fuente (color y tamaño)"""
        # Aplicar color de fuente
        if hasattr(run.font, 'color') and run.font.color.rgb:
            color = run.font.color.rgb
            hex_color = f"#{color.red:02x}{color.green:02x}{color.blue:02x}"
            text = f"<span style='color:{hex_color};'>{text}</span>"
        
        # Aplicar tamaño de fuente
        if hasattr(run.font, 'size') and run.font.size:
            font_size_pt = run.font.size.pt
            text = f"<span style='font-size:{font_size_pt:.1f}pt;'>{text}</span>"
        
        return text

    def _iter_block_items(self, parent):
        for child in parent.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _is_list_paragraph(self, paragraph):
        # Detecta si el párrafo es una lista por el estilo o por el XML
        style = paragraph.style.name.lower()
        if "list" in style or "bullet" in style:
            return True
        # XML: busca el elemento numPr
        if paragraph._element.xpath('.//w:numPr'):
            return True
        return False

    def retrieve(self, request, *args, **kwargs):
        """Obtener documento subido específico con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Documento subido obtenido exitosamente",
                code="available_field_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener el documento subido",
                code="available_field_error",
                http_status=500
            )
    
    def update(self, request, *args, **kwargs):
        """Actualizar documento subido con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            serializer.is_valid(raise_exception=True)
            serializer.save()
            
            return self.success_response(
                data=serializer.data,
                message="Documento subido actualizado exitosamente",
                code="available_field_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el documento subido",
                code="document_uploaded_update_error",
                http_status=400
            )
    
    def partial_update(self, request, *args, **kwargs):
        """Actualizar parcialmente documento subido con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            serializer.save()
            
            return self.success_response(
                data=serializer.data,
                message="Documento subido actualizado exitosamente",
                code="available_field_partial_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el documento subido",
                code="document_uploaded_update_error",
                http_status=400
            )
    
    def destroy(self, request, *args, **kwargs):
        """Eliminar documento subido con formato estándar y control de permisos"""
        try:
            instance = self.get_object()
            user = request.user
            
            # Verificar permisos para eliminar el documento
            can_delete = False
            
            # Superuser o staff pueden eliminar cualquier documento
            if user.is_superuser or user.is_staff:
                can_delete = True
            # Usuarios del grupo 'Admin' pueden eliminar documentos de su empresa
            elif user.groups.filter(name='Admin').exists():
                if user.empresa and instance.usuario.empresa == user.empresa:
                    can_delete = True
                elif not user.empresa and instance.usuario == user:
                    can_delete = True
            # Usuarios comunes solo pueden eliminar sus propios documentos
            elif instance.usuario == user:
                can_delete = True
            
            if not can_delete:
                return self.error_response(
                    errors="No tienes permisos para eliminar este documento",
                    message="Acceso denegado",
                    code="permission_denied",
                    http_status=403
                )
            
            documento_nombre = instance.nombre_original  # Guardamos el nombre antes de eliminar
            
            # El método delete() del modelo se encarga de eliminar el archivo físico automáticamente
            instance.delete()
            
            return self.success_response(
                data={"deleted_document": documento_nombre},
                message="Documento subido eliminado exitosamente (archivo físico incluido)",
                code="document_uploaded_deleted",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al eliminar el documento subido",
                code="document_uploaded_delete_error",
                http_status=500
            )

class CampoDisponibleViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = CampoDisponible.objects.all()
    serializer_class = CampoDisponibleSerializer

    def list(self, request, *args, **kwargs):
        """Listar campos disponibles con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Campos disponibles obtenidos exitosamente (paginados)",
            unpaginated_message="Campos disponibles obtenidos exitosamente",
            code="available_fields_retrieved",
            error_code="available_fields_error"
        )
    
    def create(self, request, *args, **kwargs):
        """Crear campo disponible con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Campo disponible creado exitosamente",
            code="available_field_created",
            error_message="Error al crear campo disponible",
            error_code="available_field_creation_error"
        )
    
    def retrieve(self, request, *args, **kwargs):
        """Obtener campo disponible específico con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Campo disponible obtenido exitosamente",
                code="available_field_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener el campo disponible",
                code="available_field_error",
                http_status=500
            )
    
    def update(self, request, *args, **kwargs):
        """Actualizar campo disponible con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            serializer.is_valid(raise_exception=True)
            serializer.save()
            
            return self.success_response(
                data=serializer.data,
                message="Campo disponible actualizado exitosamente",
                code="available_field_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el campo disponible",
                code="available_field_update_error",
                http_status=400
            )
    
    def partial_update(self, request, *args, **kwargs):
        """Actualizar parcialmente campo disponible con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            serializer.save()
            
            return self.success_response(
                data=serializer.data,
                message="Campo disponible actualizado exitosamente",
                code="available_field_partial_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el campo disponible",
                code="available_field_update_error",
                http_status=400
            )
    
    def destroy(self, request, *args, **kwargs):
        """Eliminar campo disponible con formato estándar"""
        try:
            instance = self.get_object()
            campo_nombre = instance.nombre  # Guardamos el nombre antes de eliminar
            instance.delete()
            return self.success_response(
                data={"deleted_field": campo_nombre},
                message="Campo disponible eliminado exitosamente",
                code="available_field_deleted",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al eliminar el campo disponible",
                code="available_field_delete_error",
                http_status=500
            )

class TipoPlantillaDocumentoViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = TipoPlantillaDocumento.objects.all()
    serializer_class = TipoPlantillaDocumentoSerializer

    def list(self, request, *args, **kwargs):
        """Listar tipos de plantilla de documento con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Tipos de plantilla de documento obtenidos exitosamente (paginados)",
            unpaginated_message="Tipos de plantilla de documento obtenidos exitosamente",
            code="tipo_plantilla_retrieved",
            error_code="tipo_plantilla_error"
        )
    
    def create(self, request, *args, **kwargs):
        """Crear tipo de plantilla de documento con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Tipo de plantilla de documento creado exitosamente",
            code="tipo_plantilla_created",
            error_message="Error al crear tipo de plantilla de documento",
            error_code="tipo_plantilla_creation_error"
        )
    
    def retrieve(self, request, *args, **kwargs):
        """Obtener tipo de plantilla de documento específico con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Tipo de plantilla de documento obtenido exitosamente",
                code="tipo_plantilla_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener el tipo de plantilla de documento",
                code="tipo_plantilla_error",
                http_status=500
            )
    
    def update(self, request, *args, **kwargs):
        """Actualizar tipo de plantilla de documento con formato estándar"""
        try:
            response = super().update(request, *args, **kwargs)
            return self.success_response(
                data=response.data,
                message="Tipo de plantilla de documento actualizado exitosamente",
                code="tipo_plantilla_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el tipo de plantilla de documento",
                code="tipo_plantilla_update_error",
                http_status=400
            )
    
    def partial_update(self, request, *args, **kwargs):
        """Actualizar parcialmente tipo de plantilla de documento con formato estándar"""
        try:
            response = super().partial_update(request, *args, **kwargs)
            return self.success_response(
                data=response.data,
                message="Tipo de plantilla de documento actualizado exitosamente",
                code="tipo_plantilla_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el tipo de plantilla de documento",
                code="tipo_plantilla_update_error",
                http_status=400
            )
    
    def destroy(self, request, *args, **kwargs):
        """Eliminar tipo de plantilla de documento con formato estándar"""
        try:
            instance = self.get_object()
            instance.delete()
            return self.success_response(
                message="Tipo de plantilla de documento eliminado exitosamente",
                code="tipo_plantilla_documento_deleted",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al eliminar el tipo de plantilla de documento",
                code="tipo_plantilla_documento_delete_error",
                http_status=500
            )

class TipoPlantillaDocumentoListAPIView(StandardResponseMixin, generics.ListAPIView):
    queryset = TipoPlantillaDocumento.objects.all()
    serializer_class = TipoPlantillaDocumentoSerializer

    def list(self, request, *args, **kwargs):
        """Listar tipos de plantilla de documento con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Tipos de plantilla de documento obtenidos exitosamente (paginados)",
            unpaginated_message="Tipos de plantilla de documento obtenidos exitosamente",
            code="tipo_plantilla_retrieved",
            error_code="tipo_plantilla_error"
        )

class CategoriaPlantillaDocumentoListAPIView(StandardResponseMixin, generics.ListAPIView):
    queryset = CategoriaPlantillaDocumento.objects.all()
    serializer_class = CategoriaPlantillaDocumentoSerializer

    def list(self, request, *args, **kwargs):
        """Listar categorias de plantilla de documento con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Categorias de plantilla de documento obtenidos exitosamente (paginados)",
            unpaginated_message="Categorias de plantilla de documento obtenidos exitosamente",
            code="categoria_plantilla_retrieved",
            error_code="categoria_plantilla_error"
        )

class PlantillaDocumentoViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = PlantillaDocumento.objects.none()
    serializer_class = PlantillaDocumentoSerializer

    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return PlantillaDocumento.objects.none()
        user = self.request.user
        # Plantillas propias o compartidas conmigo
        compartidas_ids = PlantillaCompartida.objects.filter(usuario=user).values_list('plantilla_id', flat=True)
        return PlantillaDocumento.objects.filter(
            models.Q(usuario=user) | models.Q(id__in=compartidas_ids)
        ).distinct()
    
    def list(self, request, *args, **kwargs):
        #print("list: ", request.user)
        """Listar plantillas con información de favoritos"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            usuario = request.user
            plantillas = self.get_queryset()  # Solo las del usuario
            
            # Obtener favoritos del usuario
            favoritos_usuario = PlantillaFavorita.objects.filter(usuario=usuario).values_list('plantilla_id', flat=True)
            
            plantillas_con_favoritos = []
            for plantilla in plantillas:
                # Obtener campos asociados
                campos_asociados = []
                for campo_plantilla in plantilla.campos_asociados.all():
                    campos_asociados.append({
                        'id': campo_plantilla.id,
                        'campo': campo_plantilla.campo.id,
                        'nombre_variable': campo_plantilla.nombre_variable,
                        'campo_nombre': campo_plantilla.campo.nombre,
                        'campo_tipo': campo_plantilla.campo.tipo_dato
                    })
                
                plantilla_data = {
                    'id': plantilla.id,
                    'nombre': plantilla.nombre,
                    'descripcion': plantilla.descripcion,
                    'html_con_campos': plantilla.html_con_campos,
                    'fecha_creacion': plantilla.fecha_creacion,
                    'campos_asociados': campos_asociados,
                    'es_favorito': plantilla.id in favoritos_usuario,
                    'tipo': {
                        'id': plantilla.tipo.id,
                        'nombre': plantilla.tipo.nombre
                    } if plantilla.tipo else None
                }
                plantillas_con_favoritos.append(plantilla_data)
            
            # Usar success_response directamente en lugar de paginated_list_response
            return self.success_response(
                data=plantillas_con_favoritos,
                message="Plantillas de documentos obtenidas exitosamente",
                code="plantillas_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantillas de documentos: {str(e)}",
                code="plantillas_retrieval_error"
            )

    def create(self, request, *args, **kwargs):
        """Crear plantilla de documento"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            
            serializer = self.get_serializer(data=request.data)
            if serializer.is_valid():
                plantilla = serializer.save(usuario=request.user)
                return self.standard_create_response(
                    serializer,
                    message="Plantilla de documento creada exitosamente",
                    code="plantilla_created"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para crear plantilla de documento",
                    code="plantilla_creation_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al crear plantilla de documento: {str(e)}",
                code="plantilla_creation_error"
            )

    def retrieve(self, request, *args, **kwargs):
        """Obtener plantilla de documento específica"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Plantilla de documento obtenida exitosamente",
                code="plantilla_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantilla de documento: {str(e)}",
                code="plantilla_retrieval_error"
            )

    def update(self, request, *args, **kwargs):
        """Actualizar plantilla de documento completa"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla de documento actualizada exitosamente",
                    code="plantilla_updated"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para actualizar plantilla de documento",
                    code="plantilla_update_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar plantilla de documento: {str(e)}",
                code="plantilla_update_error"
            )

    @action(detail=False, methods=['post'])
    def crear_plantilla(self, request):
        """Crear plantilla con campos asociados"""
        try:
            serializer = CrearPlantillaSerializer(data=request.data)
            if serializer.is_valid():
                # Obtener el tipo si se proporciona
                tipo = None
                if 'tipo_id' in serializer.validated_data and serializer.validated_data['tipo_id']:
                    try:
                        tipo = TipoPlantillaDocumento.objects.get(id=serializer.validated_data['tipo_id'])
                    except TipoPlantillaDocumento.DoesNotExist:
                        pass
                
                # Obtener la categoría si se proporciona
                categoria = None
                if 'categoria_id' in serializer.validated_data and serializer.validated_data['categoria_id']:
                    try:
                        categoria = CategoriaPlantillaDocumento.objects.get(id=serializer.validated_data['categoria_id'])
                    except CategoriaPlantillaDocumento.DoesNotExist:
                        pass

                # Obtener la clasificación si se proporciona
                clasificacion = None
                if 'clasificacion_id' in serializer.validated_data and serializer.validated_data['clasificacion_id']:
                    try:
                        clasificacion = ClasificacionPlantillaGeneral.objects.get(id=serializer.validated_data['clasificacion_id'])
                    except ClasificacionPlantillaGeneral.DoesNotExist:
                        pass

                # Verificar autenticación
                if not request.user.is_authenticated:
                    return self.error_response(
                        errors="Usuario no autenticado",
                        message="Autenticación requerida",
                        code="authentication_required",
                        http_status=401
                    )

                # Crear plantilla
                plantilla = PlantillaDocumento.objects.create(
                    nombre=serializer.validated_data['nombre'],
                    descripcion=serializer.validated_data.get('descripcion', ''),
                    html_con_campos=serializer.validated_data.get('html_con_campos', ''),
                    tipo=tipo,
                    categoria=categoria,
                    clasificacion=clasificacion,
                    usuario=request.user
                )

                # Crear campos asociados
                campos = serializer.validated_data.get('campos', [])
                for campo_data in campos:
                    CampoPlantilla.objects.create(
                        plantilla=plantilla,
                        campo_id=campo_data['campo_id'],
                        nombre_variable=campo_data['nombre_variable']
                    )

                return self.success_response(
                    data={'id': plantilla.id},
                    message="Plantilla creada exitosamente",
                    code="plantilla_created"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para crear plantilla",
                    code="plantilla_creation_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al crear plantilla: {str(e)}",
                code="plantilla_creation_error"
            )

    @action(detail=True, methods=['post'])
    def generar_documento(self, request, pk=None):
        """Generar documento a partir de plantilla y datos"""
        try:
            plantilla = self.get_object()
            serializer = GenerarDocumentoSerializer(data=request.data)
            
            if serializer.is_valid():
                datos = serializer.validated_data['datos']
                
                # Reemplazar campos en el HTML
                html_resultante = plantilla.html_con_campos
                for campo_plantilla in plantilla.campos_asociados.all():
                    variable = f"{{{{{campo_plantilla.nombre_variable}}}}}"
                    valor = datos.get(campo_plantilla.nombre_variable, '')
                    html_resultante = html_resultante.replace(variable, str(valor))

                # Verificar autenticación
                if not request.user.is_authenticated:
                    return self.error_response(
                        errors="Usuario no autenticado",
                        message="Autenticación requerida",
                        code="authentication_required",
                        http_status=401
                    )

                # Crear documento generado
                documento = DocumentoGenerado.objects.create(
                    plantilla=plantilla,
                    usuario=request.user,
                    datos_rellenados=datos,
                    html_resultante=html_resultante,
                    nombre=request.user.username + "_" + request.data['nombre'] + ".html"
                )

                return self.success_response(
                    data={
                        'id': documento.id,
                        'html_resultante': html_resultante
                    },
                    message="Documento generado exitosamente",
                    code="documento_generated"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para generar documento",
                    code="documento_generation_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al generar documento: {str(e)}",
                code="documento_generation_error"
            )

    def partial_update(self, request, *args, **kwargs):
        try:
            instance = self.get_object()
            data = request.data

            # Quitar campo asociado (flujo rápido)
            if 'quitar_campo_id' in data:
                campo_plantilla_id = data['quitar_campo_id']
                from .models import CampoPlantilla
                CampoPlantilla.objects.filter(id=campo_plantilla_id, plantilla=instance).delete()
                return self.success_response(
                    message="Campo asociado eliminado exitosamente",
                    code="campo_asociado_removed"
                )

            # Manejar tipo de plantilla
            if 'tipo' in data:
                tipo = None
                if data['tipo']:
                    try:
                        tipo = TipoPlantillaDocumento.objects.get(id=data['tipo'])
                    except TipoPlantillaDocumento.DoesNotExist:
                        pass
                instance.tipo = tipo
                instance.save()

            # Sincronizar campos asociados si se envía 'campos'
            if 'campos' in data:
                from .models import CampoPlantilla
                nuevos = data['campos']
                nuevos_set = set((c['campo_id'], c['nombre_variable']) for c in nuevos)
                actuales = list(instance.campos_asociados.all())
                actuales_set = set((c.campo_id, c.nombre_variable) for c in actuales)

                # Eliminar los que ya no están
                for c in actuales:
                    if (c.campo_id, c.nombre_variable) not in nuevos_set:
                        c.delete()

                # Agregar los nuevos que no existen
                for c in nuevos:
                    if (c['campo_id'], c['nombre_variable']) not in actuales_set:
                        CampoPlantilla.objects.create(
                            plantilla=instance,
                            campo_id=c['campo_id'],
                            nombre_variable=c['nombre_variable']
                        )

            # Resto del update normal
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla de documento actualizada exitosamente",
                    code="plantilla_updated"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para actualizar plantilla de documento",
                    code="plantilla_update_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar plantilla de documento: {str(e)}",
                code="plantilla_update_error"
            )

    def destroy(self, request, *args, **kwargs):
        try:
            instance = self.get_object()
            instance.delete()
            return self.success_response(
                message="Plantilla de documento eliminada exitosamente",
                code="plantilla_deleted"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al eliminar plantilla de documento: {str(e)}",
                code="plantilla_deletion_error"
            )

class DocumentoGeneradoViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = DocumentoGenerado.objects.none()
    serializer_class = DocumentoGeneradoSerializer

    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return DocumentoGenerado.objects.none()
        return DocumentoGenerado.objects.filter(usuario=self.request.user)

    def perform_create(self, serializer):
        """Asignar automáticamente el usuario logueado al crear un documento generado"""
        if not self.request.user.is_authenticated:
            raise PermissionError("Usuario no autenticado")
        serializer.save(usuario=self.request.user)
    
    def list(self, request, *args, **kwargs):
        """Listar documentos generados con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Documentos generados obtenidos exitosamente (paginados)",
            unpaginated_message="Documentos generados obtenidos exitosamente",
            code="documentos_generados_retrieved",
            error_code="documentos_generados_error"
        )
    
    def create(self, request, *args, **kwargs):
        """Crear documento generado con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Documento generado exitosamente",
            code="documento_generated",
            error_message="Error al generar documento",
            error_code="documento_generation_error"
        )
    
    def retrieve(self, request, *args, **kwargs):
        """Obtener documento generado específico con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Documento generado obtenido exitosamente",
                code="documento_generated_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener el documento generado",
                code="documento_generated_error",
                http_status=500
            )
    
    def update(self, request, *args, **kwargs):
        """Actualizar documento generado con formato estándar"""
        try:
            response = super().update(request, *args, **kwargs)
            return self.success_response(
                data=response.data,
                message="Documento generado actualizado exitosamente",
                code="documento_generated_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el documento generado",
                code="documento_generated_update_error",
                http_status=400
            )
    
    def partial_update(self, request, *args, **kwargs):
        """Actualizar parcialmente documento generado con formato estándar"""
        try:
            response = super().partial_update(request, *args, **kwargs)
            return self.success_response(
                data=response.data,
                message="Documento generado actualizado exitosamente",
                code="documento_generated_updated",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al actualizar el documento generado",
                code="documento_generated_update_error",
                http_status=400
            )

    def destroy(self, request, *args, **kwargs):
        """Eliminar documento generado con formato estándar"""
        try:
            instance = self.get_object()
            campo_nombre = instance.nombre  # Guardamos el nombre antes de eliminar
            instance.delete()
            return self.success_response(
                data={"deleted_generated_document": campo_nombre},
                message="Documento generado eliminado exitosamente",
                code="documento_generated_deleted",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al eliminar el documento generado",
                code="documento_generated_delete_error",
                http_status=500
            )

class PlantillaFavoritaViewSet(StandardResponseMixin, viewsets.GenericViewSet):
    queryset = PlantillaFavorita.objects.none()
    serializer_class = PlantillaFavoritaSerializer
    '''
    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return PlantillaFavorita.objects.none()
        return PlantillaFavorita.objects.filter(usuario=self.request.user)
    
    def list(self, request, *args, **kwargs):
        """Listar favoritos del usuario con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Favoritos obtenidos exitosamente (paginados)",
            unpaginated_message="Favoritos obtenidos exitosamente",
            code="favorites_retrieved",
            error_code="favorites_error"
        )
    
    def create(self, request, *args, **kwargs):
        """Crear favorito con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Favorito creado exitosamente",
            code="favorite_created",
            error_message="Error al crear favorito",
            error_code="favorite_create_error"
        )
    '''
    @action(detail=False, methods=['post'])
    def agregar_favorito(self, request):
        """Agregar plantilla a favoritos"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            
            plantilla_id = request.data.get('plantilla_id')
            if not plantilla_id:
                return self.error_response(
                    errors="plantilla_id es requerido",
                    message="Datos incompletos",
                    code="missing_data",
                    http_status=400
                )
            
            plantilla = PlantillaDocumento.objects.get(id=plantilla_id)
            usuario = request.user
            
            # Verificar si ya existe
            favorito, created = PlantillaFavorita.objects.get_or_create(
                usuario=usuario,
                plantilla=plantilla
            )
            
            if created:
                serializer = self.serializer_class(favorito)
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla agregada a favoritos exitosamente",
                    code="plantilla_favorita_added",
                    http_status=201
                )
            else:
                serializer = self.serializer_class(favorito)
                return self.success_response(
                    data=serializer.data,
                    message="La plantilla ya está en favoritos",
                    code="plantilla_favorita_exists",
                    http_status=200
                )
                
        except PlantillaDocumento.DoesNotExist:
            return self.error_response(
                errors="Plantilla no encontrada",
                message="Recurso no encontrado",
                code="plantilla_not_found",
                http_status=404
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error interno del servidor",
                code="server_error",
                http_status=500
            )
    
    @action(detail=False, methods=['delete'])
    def quitar_favorito(self, request):
        """Quitar plantilla de favoritos"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            
            plantilla_id = request.data.get('plantilla_id')
            if not plantilla_id:
                return self.error_response(
                    errors="plantilla_id es requerido",
                    message="Datos incompletos",
                    code="missing_data",
                    http_status=400
                )
            
            usuario = request.user
            
            favorito = PlantillaFavorita.objects.get(
                usuario=usuario,
                plantilla_id=plantilla_id
            )
            favorito.delete()
            
            return self.success_response(
                data=None,
                message="Plantilla removida de favoritos exitosamente",
                code="plantilla_favorita_removed",
                http_status=200
            )
                
        except PlantillaFavorita.DoesNotExist:
            return self.error_response(
                errors="La plantilla no está en favoritos",
                message="Recurso no encontrado",
                code="not_found",
                http_status=404
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error interno del servidor",
                code="server_error",
                http_status=500
            )
    
    @action(detail=False, methods=['get'])
    def mis_favoritos(self, request):
        """Obtener plantillas favoritas del usuario"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    errors="Usuario no autenticado",
                    message="Autenticación requerida",
                    code="authentication_required",
                    http_status=401
                )
            
            usuario = request.user
            favoritos = PlantillaFavorita.objects.filter(usuario=usuario).select_related('plantilla')
            
            plantillas_favoritas = []
            for favorito in favoritos:
                plantilla = favorito.plantilla
                plantillas_favoritas.append({
                    'id': plantilla.id,
                    'nombre': plantilla.nombre,
                    'descripcion': plantilla.descripcion,
                    'fecha_creacion': plantilla.fecha_creacion,
                    'fecha_agregado_favorito': favorito.fecha_agregado,
                    'es_favorito': True
                })
            
            return self.success_response(
                data=plantillas_favoritas,
                message="Plantillas favoritas obtenidas exitosamente",
                code="plantilla_favorita_retrieved",
                http_status=200
            )
                
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener favoritos",
                code="plantilla_favorita_error",
                http_status=500
            )

class PlantillaCompartidaViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    queryset = PlantillaCompartida.objects.none()
    serializer_class = PlantillaCompartidaSerializer

    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return PlantillaCompartida.objects.none()
        user = self.request.user
        return PlantillaCompartida.objects.filter(models.Q(usuario=user) | models.Q(plantilla__usuario=user)).distinct()

    def list(self, request, *args, **kwargs):
        """Listar plantillas compartidas con formato estándar"""
        return self.paginated_list_response(
            request=request,
            queryset=self.get_queryset(),
            serializer_class=self.serializer_class,
            paginated_message="Plantillas compartidas obtenidas exitosamente (paginadas)",
            unpaginated_message="Plantillas compartidas obtenidas exitosamente",
            code="plantilla_compartida_retrieved",
            error_code="plantilla_compartida_error"
        )

    def create(self, request, *args, **kwargs):
        """Crear plantilla compartida con formato estándar"""
        return self.standard_create_response(
            request=request,
            success_message="Plantilla compartida exitosamente",
            code="plantilla_compartida",
            error_message="Error al compartir plantilla",
            error_code="plantilla_compartida_error"
        )

    def retrieve(self, request, *args, **kwargs):
        """Obtener plantilla compartida específica con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Plantilla compartida obtenida exitosamente",
                code="plantilla_compartida_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantilla compartida: {str(e)}",
                code="plantilla_compartida_retrieval_error"
            )

    def update(self, request, *args, **kwargs):
        """Actualizar plantilla compartida con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla compartida actualizada exitosamente",
                    code="plantilla_compartida_updated"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para actualizar plantilla compartida",
                    code="plantilla_compartida_update_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar plantilla compartida: {str(e)}",
                code="plantilla_compartida_update_error"
            )

    def partial_update(self, request, *args, **kwargs):
        """Actualizar parcialmente plantilla compartida con formato estándar"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla compartida actualizada exitosamente",
                    code="plantilla_compartida_updated"
                )
            else:
                return self.error_response(
                    message="Datos inválidos para actualizar plantilla compartida",
                    code="plantilla_compartida_update_error",
                    errors=serializer.errors
                )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar plantilla compartida: {str(e)}",
                code="plantilla_compartida_update_error"
            )

    def destroy(self, request, *args, **kwargs):
        """Eliminar plantilla compartida con formato estándar"""
        try:
            instance = self.get_object()
            instance.delete()
            return self.success_response(
                message="Plantilla compartida eliminada exitosamente",
                code="plantilla_compartida_deleted"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al eliminar plantilla compartida: {str(e)}",
                code="plantilla_compartida_deletion_error"
            )

    @action(detail=False, methods=['get'])
    def compartidas_conmigo(self, request):
        """Obtener plantillas compartidas conmigo con formato estándar"""
        try:
            user = request.user
            if not user or not user.is_authenticated:
                return self.error_response(
                    message="Usuario no autenticado",
                    code="authentication_required"
                )
            
            compartidas = PlantillaCompartida.objects.filter(usuario=user)
            serializer = self.get_serializer(compartidas, many=True)
            return self.success_response(
                data=serializer.data,
                message="Plantillas compartidas conmigo obtenidas exitosamente",
                code="compartida_conmigo_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantillas compartidas conmigo: {str(e)}",
                code="compartida_conmigo_error"
            )

    @action(detail=False, methods=['post'])
    def compartir(self, request):
        """Compartir una plantilla con uno o varios usuarios con formato estándar"""
        try:
            plantilla_id = request.data.get('plantilla_id')
            usuario_id = request.data.get('usuario_id')
            usuario_ids = request.data.get('usuario_ids')
            permisos = request.data.get('permisos', 'lectura')
            
            if not plantilla_id or (not usuario_id and not usuario_ids):
                return self.error_response(
                    message="plantilla_id y usuario_id(s) son requeridos",
                    code="missing_required_fields"
                )
            
            ids = []
            if usuario_ids:
                ids = usuario_ids if isinstance(usuario_ids, list) else [usuario_ids]
            elif usuario_id:
                ids = [usuario_id]
            
            compartidas = []
            for uid in ids:
                compartida, created = PlantillaCompartida.objects.get_or_create(
                    plantilla_id=plantilla_id,
                    usuario_id=uid,
                    defaults={'permisos': permisos}
                )
                if not created:
                    compartida.permisos = permisos
                    compartida.save()
                compartidas.append({'id': compartida.id, 'usuario_id': uid})
            
            return self.success_response(
                data={'compartidas': compartidas},
                message="Plantilla compartida correctamente",
                code="plantilla_compartida_successfully"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al compartir plantilla: {str(e)}",
                code="plantilla_compartida_error"
            )

    @action(detail=True, methods=['delete'])
    def revocar(self, request, pk=None):
        """Revocar acceso a una plantilla compartida con formato estándar"""
        try:
            compartida = self.get_object()
            compartida.delete()
            return self.success_response(
                message="Acceso revocado exitosamente",
                code="access_revoked"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al revocar acceso: {str(e)}",
                code="access_revoke_error"
            )

    @action(detail=False, methods=['get'])
    def usuarios_compartidos(self, request):
        """Obtener usuarios con los que se ha compartido una plantilla con formato estándar"""
        try:
            plantilla_id = request.query_params.get('plantilla_id')
            if not plantilla_id:
                return self.error_response(
                    message="plantilla_id es requerido",
                    code="missing_plantilla_id"
                )
            
            compartidas = PlantillaCompartida.objects.filter(plantilla_id=plantilla_id)
            data = [
                {
                    'id': c.id,
                    'usuario': c.usuario.id,
                    'usuario_username': c.usuario.username,
                    'usuario_email': c.usuario.email,
                    'permisos': c.permisos,
                    'fecha_compartida': c.fecha_compartida,
                }
                for c in compartidas
            ]
            return self.success_response(
                data=data,
                message="Usuarios compartidos obtenidos exitosamente",
                code="shared_users_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener usuarios compartidos: {str(e)}",
                code="shared_users_error"
            )

class ClasificacionPlantillaGeneralViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = ClasificacionPlantillaGeneral.objects.all()
    serializer_class = ClasificacionPlantillaGeneralSerializer

    def list(self, request, *args, **kwargs):
        """Listar clasificaciones de plantilla general"""
        try:
            queryset = self.filter_queryset(self.get_queryset())
            return self.paginated_list_response(
                request,
                queryset,
                self.get_serializer_class(),
                paginated_message="Listado paginado de clasificaciones de plantilla general obtenido correctamente",
                unpaginated_message="Listado de clasificaciones de plantilla general obtenido correctamente",
                code="classification_list_retrieved",
                error_code="classification_list_error"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener las clasificaciones de plantilla general: {str(e)}",
                code="classification_list_error"
            )

    def create(self, request, *args, **kwargs):
        """Crear nueva clasificación de plantilla general"""
        try:
            return self.standard_create_response(
                request,
                self.get_serializer_class(),
                success_message="Clasificación de plantilla general creada correctamente",
                success_code="classification_created",
                error_message="Error al crear la clasificación de plantilla general",
                error_code="classification_creation_error"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al crear la clasificación de plantilla general: {str(e)}",
                code="classification_creation_error"
            )

    def retrieve(self, request, *args, **kwargs):
        """Obtener una clasificación de plantilla general específica"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Clasificación de plantilla general obtenida correctamente",
                code="classification_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener la clasificación de plantilla general: {str(e)}",
                code="classification_retrieval_error"
            )

    def update(self, request, *args, **kwargs):
        """Actualizar clasificación de plantilla general completa"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Clasificación de plantilla general actualizada correctamente",
                    code="classification_updated"
                )
            return self.error_response(
                message="Datos inválidos para actualizar la clasificación",
                code="classification_update_validation_error",
                data=serializer.errors
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar la clasificación de plantilla general: {str(e)}",
                code="classification_update_error"
            )

    def partial_update(self, request, *args, **kwargs):
        """Actualizar clasificación de plantilla general parcial"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Clasificación de plantilla general actualizada correctamente",
                    code="classification_updated"
                )
            return self.error_response(
                message="Datos inválidos para actualizar la clasificación",
                code="classification_update_validation_error",
                data=serializer.errors
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar la clasificación de plantilla general: {str(e)}",
                code="classification_update_error"
            )

    def destroy(self, request, *args, **kwargs):
        """Eliminar clasificación de plantilla general"""
        try:
            instance = self.get_object()
            instance.delete()
            return self.success_response(
                message="Clasificación de plantilla general eliminada correctamente",
                code="classification_deleted"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al eliminar la clasificación de plantilla general: {str(e)}",
                code="classification_deletion_error"
            )

class ClasificacionPlantillaGeneralListAPIView(StandardResponseMixin, generics.ListAPIView):
    queryset = ClasificacionPlantillaGeneral.objects.all()
    serializer_class = ClasificacionPlantillaGeneralSerializer

    def list(self, request, *args, **kwargs):
        """Listar clasificaciones de plantilla general"""
        try:
            queryset = self.filter_queryset(self.get_queryset())
            return self.paginated_list_response(
                request,
                queryset,
                self.get_serializer_class(),
                paginated_message="Listado paginado de clasificaciones de plantilla general obtenido correctamente",
                unpaginated_message="Listado de clasificaciones de plantilla general obtenido correctamente",
                code="classification_list_retrieved",
                error_code="classification_list_error"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener las clasificaciones de plantilla general: {str(e)}",
                code="classification_list_error"
            )

class PlantillaGeneralViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    queryset = PlantillaGeneral.objects.all()
    serializer_class = PlantillaGeneralSerializer

    def list(self, request, *args, **kwargs):
        """Listar plantillas generales"""
        try:
            queryset = self.filter_queryset(self.get_queryset())
            return self.paginated_list_response(
                request,
                queryset,
                self.get_serializer_class(),
                paginated_message="Listado paginado de plantillas generales obtenido correctamente",
                unpaginated_message="Listado de plantillas generales obtenido correctamente",
                code="plantilla_general_retrieved",
                error_code="plantilla_general_list_error"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener las plantillas generales: {str(e)}",
                code="plantilla_general_list_error"
            )

    def create(self, request, *args, **kwargs):
        """Crear nueva plantilla general"""
        try:
            return self.standard_create_response(
                request,
                self.get_serializer_class(),
                success_message="Plantilla general creada correctamente",
                success_code="plantilla_general_created",
                error_message="Error al crear la plantilla general",
                error_code="plantilla_general_creation_error"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al crear la plantilla general: {str(e)}",
                code="plantilla_general_creation_error"
            )

    def retrieve(self, request, *args, **kwargs):
        """Obtener una plantilla general específica"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance)
            return self.success_response(
                data=serializer.data,
                message="Plantilla general obtenida correctamente",
                code="plantilla_general_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener la plantilla general: {str(e)}",
                code="plantilla_general_retrieval_error"
            )

    def update(self, request, *args, **kwargs):
        """Actualizar plantilla general completa"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla general actualizada correctamente",
                    code="plantilla_general_updated"
                )
            return self.error_response(
                message="Datos inválidos para actualizar la plantilla",
                code="plantilla_general_update_validation_error",
                data=serializer.errors
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar la plantilla general: {str(e)}",
                code="plantilla_general_update_error"
            )

    def partial_update(self, request, *args, **kwargs):
        """Actualizar plantilla general parcial"""
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return self.success_response(
                    data=serializer.data,
                    message="Plantilla general actualizada correctamente",
                    code="plantilla_general_updated"
                )
            return self.error_response(
                message="Datos inválidos para actualizar la plantilla",
                code="plantilla_general_update_validation_error",
                data=serializer.errors
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al actualizar la plantilla general: {str(e)}",
                code="plantilla_general_update_error"
            )

    def destroy(self, request, *args, **kwargs):
        """Eliminar plantilla general"""
        try:
            instance = self.get_object()
            instance.delete()
            return self.success_response(
                message="Plantilla general eliminada correctamente",
                code="plantilla_general_deleted"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al eliminar la plantilla general: {str(e)}",
                code="plantilla_general_deletion_error"
            )

    @action(detail=True, methods=['get'])
    def plantillas_por_clasificacion(self, request, pk=None):
        """Obtener plantillas agrupadas por clasificación"""
        try:
            plantilla_general = self.get_object()
            data = plantilla_general.get_plantillas_por_clasificacion()
            return self.success_response(
                data=data,
                message="Plantillas por clasificación obtenidas correctamente",
                code="plantilla_general_by_classification_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantillas por clasificación: {str(e)}",
                code="plantilla_general_by_classification_error"
            )

    @action(detail=True, methods=['post'])
    def compartir(self, request, pk=None):
        """Compartir plantilla general con usuarios específicos"""
        try:
            plantilla_general = self.get_object()
            usuarios_ids = request.data.get('usuarios_ids', [])
            fecha_vencimiento = request.data.get('fecha_vencimiento')
            
            if not usuarios_ids:
                return self.error_response(
                    message="Debe especificar al menos un usuario",
                    code="no_users_specified"
                )
            
            # Crear registros de compartición
            compartidas_creadas = []
            for usuario_id in usuarios_ids:
                try:
                    usuario = User.objects.get(id=usuario_id)
                    compartida, created = PlantillaGeneralCompartida.objects.get_or_create(
                        plantilla_general=plantilla_general,
                        usuario_asignado=usuario,
                        defaults={
                            'usuario_que_comparte': request.user,
                            'fecha_vencimiento': fecha_vencimiento
                        }
                    )
                    if created:
                        compartidas_creadas.append(compartida)
                except Usuario.DoesNotExist:
                    continue
            
            return self.success_response(
                data={
                    'compartidas_creadas': len(compartidas_creadas),
                    'total_usuarios': len(usuarios_ids)
                },
                message=f"Plantilla compartida con {len(compartidas_creadas)} usuarios",
                code="plantilla_general_shared"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al compartir plantilla: {str(e)}",
                code="plantilla_general_sharing_error"
            )

    @action(detail=True, methods=['get'])
    def usuarios_con_acceso(self, request, pk=None):
        """Obtener usuarios que tienen acceso a esta plantilla general"""
        try:
            plantilla_general = self.get_object()
            compartidas = PlantillaGeneralCompartida.objects.filter(
                plantilla_general=plantilla_general
            ).select_related('usuario', 'asignado_por')
            
            usuarios_data = []
            for compartida in compartidas:
                usuarios_data.append({
                    'id': compartida.usuario.id,
                    'username': compartida.usuario.username,
                    'email': compartida.usuario.email,
                    'fecha_asignacion': compartida.fecha_asignacion,
                    'fecha_expiracion': compartida.fecha_expiracion,
                    'esta_vigente': compartida.esta_vigente(),
                    'asignado_por': compartida.asignado_por.username
                })
            
            return self.success_response(
                data=usuarios_data,
                message="Usuarios con acceso obtenidos correctamente",
                code="plantilla_general_usuarios_con_acceso_retrieved"
            )
        except Exception as e:
            return self.error_response(
                 message=f"Error al obtener usuarios con acceso: {str(e)}",
                 code="plantilla_general_usuarios_con_acceso_error"
             )


class PlantillaGeneralCompartidaViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    """ViewSet para gestionar plantillas generales compartidas"""
    queryset = PlantillaGeneralCompartida.objects.all()
    serializer_class = PlantillaGeneralCompartidaSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['plantilla_general__nombre', 'usuario__username', 'asignado_por__username']
    ordering_fields = ['fecha_asignacion', 'fecha_expiracion']
    ordering = ['-fecha_asignacion']

    def get_queryset(self):
        """Filtrar plantillas compartidas según el usuario"""
        user = self.request.user
        if not user.is_authenticated:
            return PlantillaGeneralCompartida.objects.none()
        
        if user.is_staff:
            # Los administradores pueden ver todas las plantillas compartidas
            return PlantillaGeneralCompartida.objects.all().select_related(
                'plantilla_general', 'usuario', 'asignado_por'
            )
        else:
            # Los usuarios solo pueden ver las plantillas compartidas con ellos
            return PlantillaGeneralCompartida.objects.filter(
                usuario=user
            ).select_related(
                'plantilla_general', 'usuario', 'asignado_por'
            )

    @action(detail=False, methods=['get'])
    def mis_plantillas_compartidas(self, request):
        """Obtener plantillas compartidas con el usuario actual"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    message="Usuario no autenticado",
                    code="authentication_required",
                    http_status=401
                )
            
            compartidas = self.get_queryset().filter(usuario=request.user)
            serializer = self.get_serializer(compartidas, many=True)
            return self.success_response(
                data=serializer.data,
                message="Plantillas compartidas obtenidas correctamente",
                code="plantilla_general_compartida_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantillas compartidas: {str(e)}",
                code="plantilla_general_compartida_error"
            )

    @action(detail=False, methods=['get'])
    def plantillas_vigentes(self, request):
        """Obtener solo las plantillas compartidas vigentes"""
        try:
            if not request.user.is_authenticated:
                return self.error_response(
                    message="Usuario no autenticado",
                    code="authentication_required",
                    http_status=401
                )
            
            compartidas = self.get_queryset().filter(usuario=request.user)
            vigentes = [c for c in compartidas if c.esta_vigente()]
            serializer = self.get_serializer(vigentes, many=True)
            return self.success_response(
                data=serializer.data,
                message="Plantillas vigentes obtenidas correctamente",
                code="plantilla_general_compartida_vigente_retrieved"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al obtener plantillas vigentes: {str(e)}",
                code="plantilla_general_compartida_vigente_error"
            )

    @action(detail=True, methods=['delete'])
    def revocar_acceso(self, request, pk=None):
        """Revocar acceso a una plantilla compartida"""
        try:
            compartida = self.get_object()
            # Solo el usuario que asignó o un admin puede revocar el acceso
            if request.user != compartida.asignado_por and not request.user.is_staff:
                return self.error_response(
                    message="No tiene permisos para revocar este acceso",
                    code="permission_denied"
                )
            
            compartida.delete()
            return self.success_response(
                message="Acceso revocado correctamente",
                code="access_revoked"
            )
        except Exception as e:
            return self.error_response(
                message=f"Error al revocar acceso: {str(e)}",
                code="revoke_access_error"
            )



class UsuariosViewSet(StandardResponseMixin, viewsets.ModelViewSet):
    """ViewSet para listar usuarios (excluyendo al usuario actual)"""
    queryset = Usuarios.objects.none()
    serializer_class = UsuariosSerializer
    
    def get_queryset(self):
        """Obtener todos los usuarios excepto el usuario actual"""
        if self.request.user.is_authenticated:
            return Usuarios.objects.exclude(id=self.request.user.id)
        return Usuarios.objects.none()
    
    def list(self, request, *args, **kwargs):
        """Listar usuarios con formato estándar"""
        try:
            queryset = self.get_queryset()
            serializer = self.get_serializer(queryset, many=True)
            return self.success_response(
                data=serializer.data,
                message="Usuarios obtenidos exitosamente",
                code="users_retrieved",
                http_status=200
            )
        except Exception as e:
            return self.error_response(
                errors=str(e),
                message="Error al obtener usuarios",
                code="users_error",
                http_status=500
            )